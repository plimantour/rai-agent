
# Philippe Limantour - March 2024
# This file contains the functions to update a word document

from pdfminer.high_level import extract_text
import docx2txt
import docx
import os
import re
from termcolor import colored
import time
import hashlib

try:
    from termcolor import colored
except ImportError:
    def colored(x, *args, **kwargs):
        return x

# Method to generate a unique identifier
def generate_unique_identifier():
    # Generate the unique identifier using MD5 hash to allow multiple users to generate RAI Impact Assessments at the same time
    timestamp = time.strftime("%Y%m%d%H%M%S")
    identifier = hashlib.md5(timestamp.encode('utf-8')).hexdigest()
    return identifier

# Method to get an input_text and a word file name, and save the text to the word file
def save_text_to_docx(input_text, docx_filename_path):
    try:
        folder = os.path.dirname(docx_filename_path)
        if not os.path.exists(folder):
            os.makedirs(folder)
        doc = docx.Document()
        doc.add_paragraph(input_text)
        doc.save(docx_filename_path)
        return True
    except Exception as e:
        print(colored(f"Error saving text to file {docx_filename_path}:\n{e}", "red"))
        return False

# Method to get the filename without extension and the extension from a file path
def get_filename_and_extension(file_path):
    # Get the base name of the file
    base_name = os.path.basename(file_path)
    # Split the base name into name and extension
    file_name, file_extension = os.path.splitext(base_name)
    return file_name, file_extension

# Decorator to display the execution time of a function: @timer_decorator above a def funtion(...)
def timer_decorator(func):
    def wrapper(*args, **kwargs):
        start_time = time.time()
        result = func(*args, **kwargs)
        end_time = time.time()
        exec_time = end_time - start_time
        print(colored(f"Execution time for {func.__name__}: {exec_time:.4f} seconds", "cyan"))
        return result
    return wrapper

def extract_text_from_pdf(file):
    return extract_text(file)

def extract_text_from_docx(file):
    return docx2txt.process(file)

# Method to extract text from a file - pdf, docx, or txt
# Returns filename without extension and extracted text
def extract_text_from_input(file_path):
    if not os.path.exists(file_path):
        print(colored(f"File {file_path} does not exist", "red"))
        return None, None

    filename = os.path.basename(file_path)
    filename_without_extension = os.path.splitext(filename)[0]
    extension = os.path.splitext(filename)[1]

    if extension == ".pdf":
        text = extract_text_from_pdf(file_path)
        return filename_without_extension, text
    elif extension == ".docx":
        text = extract_text_from_docx(file_path)
        return filename_without_extension, text
    elif extension == '.txt' or extension == '.json':
        try:
            with open(file_path, "r") as file:
                return filename_without_extension, file.read()
        except Exception as e:
            print(colored(f"Error reading file {file_path}:\n{e}", "red"))
            return None, None
    else:
        print(colored(f"Unsupported file type: {extension} for input {file_path}", "red"))
        return None, None

# Method to extract text from an uploaded file
def extract_text_from_upload(file):
    if file.type == "application/pdf":
        text = extract_text_from_pdf(file)
        # text = extract_text_from_PyMuPDF(file)
        return text, file.name
    elif (
        file.type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        text = extract_text_from_docx(file)
        return text, file.name
    elif file.type == "application/json":
        return file.getvalue().decode("utf-8"), file.name
    else:
        return file.getvalue().decode("utf-8"), file.name

@timer_decorator
def docx_find_replace_text_bydict(docx_filepath, search_replace_dict={}, search_prefix='##', doc=None, verbose=False):

    def search_text_in_paragraph(p, search_prefix, search_replace_dict):
        # Get the text from search_prefix_index until either space, carriage return, or end of text
        try:
            extracted_text = None
            replace_by_text = None
            matches = re.findall(r'##\w+', p.text)
            if len(matches) > 0:
                extracted_text = matches[0]
                replace_by_text = search_replace_dict.get(extracted_text)
                if replace_by_text is not None:
                    replace_by_text = replace_by_text.strip()
                    del search_replace_dict[extracted_text]
                    return extracted_text, replace_by_text
                else:
                    if verbose:
                        print(colored(f"Text not found in the dictionary: '{extracted_text}'", "yellow"))
                    return None, None
            else:
                return None, None
        except Exception as e:
            print(colored(f"Error extracting text from paragraph:\n{e}", "red"))
            return None, None
        
    verbose = False
        
    try:
        # Open the docx file if not provided with an already opened docx object - saving time when doing multiple replacements
        if not doc:
            doc = docx.Document(docx_filepath)

        paragraphs = list(doc.paragraphs)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraphs.append(paragraph)

        if verbose:
            print(f'Searching for <{search_replace_dict}>')

        for p in paragraphs:
            if search_prefix in p.text:
                if p.text.find(search_prefix) != -1:
                    inline = p.runs
                    found_in_sub_run = False
                    for i in range(len(inline)):
                        if search_prefix in inline[i].text:
                            extracted_sub_text, replace_by_sub_text = search_text_in_paragraph(inline[i], search_prefix, search_replace_dict)
                            if extracted_sub_text is not None and replace_by_sub_text is not None:
                                if verbose:
                                    print(f'Sub-Replacing <{extracted_sub_text}> with <{replace_by_sub_text}>')
                                inline[i].text = inline[i].text.replace(extracted_sub_text, replace_by_sub_text)
                                found_in_sub_run = True
                    if not found_in_sub_run:
                        extracted_text, replace_by_text = search_text_in_paragraph(p, search_prefix, search_replace_dict)
                        if verbose:
                            print(f'Replacing <{extracted_text}> with <{replace_by_text}>')
                        if extracted_text is not None and replace_by_text is not None:
                            p.text = p.text.replace(extracted_text, replace_by_text)

        doc.save(docx_filepath)
        return doc

    except Exception as e:
        print(colored(f"Error replacing text in file {docx_filepath}:\n{e}", "red"))
        return None


# Method to search and replace text in a docx file
@timer_decorator
def docx_find_replace_text(docx_filepath, search_text_list=[], replace_text_list=[], doc=None, one_pass=True, verbose=False):
    try:
        # Open the docx file if not provided with an already opened docx object - saving time when doing multiple replacements
        if not doc:
            doc = docx.Document(docx_filepath)

        paragraphs = list(doc.paragraphs)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraphs.append(paragraph)

        if verbose:
            print(f'Searching for <{search_text_list}>')

        for p in paragraphs:
            for index, (search_text, replace_text) in enumerate(zip(search_text_list, replace_text_list)):
                if search_text in p.text:
                    if verbose:
                        print(f'Replacing <{search_text}> with <{replace_text}>')
                    inline = p.runs
                    found_in_sub_run = False
                    for i in range(len(inline)):
                        if search_text in inline[i].text:
                            inline[i].text = inline[i].text.replace(search_text, replace_text)
                            found_in_sub_run = True
                    if not found_in_sub_run:
                        p.text = p.text.replace(search_text, replace_text)
                    if one_pass:    # Speed up the process by removing the search_text and replace_text from the lists
                        search_text_list.pop(index)    # Remove the current search_text from the list - we don't want to replace it again
                        replace_text_list.pop(index)  # Remove the corresponding replace_text from the list - only 1 occurrence for this context
                        # search_text_list.remove(search_text)    # Remove the search_text from the list - we don't want to replace it again
                        # replace_text_list.remove(replace_text)  # Remove the replace_text from the list - only 1 occurrence for this context
                        break   # We found the search_text in the paragraph, no need to continue searching
                                # (this is a simplification, it may not be the best approach for all cases but works in this context)
                                # If you need to replace the same search_text with different replace_text in the same paragraph, set one_pass to False

        doc.save(docx_filepath)

        if verbose:
            if len(search_text_list) > 0:
                print(colored(f"Search text not found in the document: {search_text_list}", "yellow"))

        return doc
    except Exception as e:
        print(colored(f"Error replacing text in file {docx_filepath}:\n{e}", "red"))
        return None

# Method to delete all text between two searched texts in a docx file - Deletes first search text, not the stop one
def docx_delete_all_between_searched_texts(docx_filepath, start_text, stop_text, doc=None, verbose=False):

    def delete_table(table):
        try:
            tbl = table._tbl  # get xml element that represents the table
            tbl.getparent().remove(tbl)  # remove table from its parent element
        except Exception as e:
            print(colored(f"Error deleting table:\n{e}", "red"))

    try:
        # Open the docx file if not provided with an already opened docx object - saving time when doing multiple replacements
        if not doc:
            doc = docx.Document(docx_filepath)
        
        verbose = False

        if verbose:
            print(f'start_text: {start_text} - stop_text: {stop_text}')

        start_index = False
        stop_index = False
        element_index = 0  # counter for elements
        para_index = 0  # counter for paragraphs
        para_start_index = -1   # index of the paragraph containing the start_text
        para_stop_index = -1     # index of the paragraph containing the stop_text
        table_index = 0  # counter for tables
        para_index_shift = 0  # counter for paragraphs to adjust the index when deleting paragraphs
        table_index_shift = 0  # counter for tables to adjust the index when deleting tables
        for element in doc.element.body:
            if element.tag.endswith('}p'):
                # This is a paragraph
                para = doc.paragraphs[para_index]
                # print('Paragraph:', para.text)
                if start_text in para.text:
                    start_index = True
                    para_start_index = para_index
                    if verbose:
                        print(f'Start index: {para_index}')
                if stop_text in para.text:
                    stop_index = True
                    para_stop_index = para_index
                    if verbose:
                        print(f'Stop index: {para_index}')
                    break
                para_index += 1  # increment the counter
                if start_index and not stop_index:
                    para.clear()
                    if verbose:
                        print(f'Deleting paragraph at index {para_index}')
                    # delete_paragraph(para)
                    # doc.paragraphs.pop(para_index)
                    # p = doc.paragraphs[para_index - para_index_shift]._element
                    # p.getparent().remove(p)
                    # p._p = p._element = None
                    para_index_shift += 1
            elif element.tag.endswith('}tbl'):
                # This is a table
                if start_index and not stop_index:
                    if verbose:
                        print(f'Deleting table at index {table_index - table_index_shift} / {len(doc.tables)} - {table_index_shift}')
                    delete_table(doc.tables[table_index - table_index_shift])
                    table_index_shift += 1
                table_index += 1
            elif element.tag.endswith('}sectPr'):
                # This is the end of the document
                break
            element_index += 1

        # Remove empty paragraphs
        p_index_shift = 0
        for p_index, paragraph in enumerate(doc.paragraphs):
            if p_index >= para_start_index and p_index <= para_stop_index and not paragraph.text.strip():  # if the paragraph is empty or contains only spaces
                p = paragraph._element
                p.getparent().remove(p)
                p._p = p._element = None
                if verbose:
                    print(f'Deleting empty paragraph at index {p_index - p_index_shift} - {p_index} / {len(doc.paragraphs)}')
                if p_index - p_index_shift < len(doc.paragraphs):
                    doc.paragraphs.pop(p_index - p_index_shift)
                    p_index_shift += 1

        doc.save(docx_filepath)

        return doc
    
    except Exception as e:
        print(colored(f"Error deleting text in file {docx_filepath}:\n{e}", "red"))
        return None

# Method to read the content of a log file
def read_logs_content(logs_filename_path="./logs.txt"):
    if not os.path.exists(logs_filename_path):
        print(colored(f"No logs file {logs_filename_path}", "red"))
        return 'Logs file not found'
    else:
        try:
            with open(logs_filename_path, 'r') as f:
                logs = f.read()
            return logs
        except Exception as e:
            print(colored(f"Error reading logs from file {logs_filename_path}:\n{e}", "red"))
            return 'Error reading logs'

# Method to append a log to a file
def append_log_to_file(log, filename_path="./logs.txt"):
    if not os.path.exists(filename_path):
        print(colored(f"No log file {filename_path}:\n{e}", "red"))
    else:
        try:
            timestamp = time.strftime("%Y%m%d%H%M%S")
            log = f'{timestamp} - {log}'
            with open(filename_path, 'a') as f:
                f.write(log)
                f.write('\n')
        except Exception as e:
            print(colored(f"Error appending log {log} to file {filename_path}:\n{e}", "red"))

# Method to get the users from a file
def get_users_list(users_filename_path="./users.txt"):
    if not os.path.exists(users_filename_path):
        print(colored(f"No users file {users_filename_path}", "red"))
        return []
    else:
        try:
            with open(users_filename_path, 'r') as f:
                users = f.readlines()
                users = [user.strip() for user in users]
            return users
        except Exception as e:
            print(colored(f"Error reading users from file {users_filename_path}:\n{e}", "red"))
            return None